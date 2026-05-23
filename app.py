import streamlit as st
import json
import os
import time
import smtplib
import io
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from groq import Groq
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.platypus import Paragraph, Table, TableStyle
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_LEFT
from PIL import Image
from pypdf import PdfReader, PdfWriter
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import requests
from lxml import etree

# ==========================================
# API KEY & MAIL CONFIG
# ==========================================
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GMAIL_USER = os.getenv("GMAIL_USER", "")
GMAIL_PASSWORD = os.getenv("GMAIL_PASSWORD", "")

# ==========================================
# SPACING & MARGINS
# ==========================================
SPACING = {
    'section_gap': 18,
    'heading_gap': 10,
    'bullet_gap': 5,
    'paragraph_gap': 12,
    'table_row_gap': 4,
}

MARGINS = {
    'left': 50,
    'right': 50,
    'top': 15,
    'bottom': 50,
    'page_border': 10,
}

# ==========================================
# CAREER TABLE DATA
# ==========================================
CAREER_TEMPLATES = {
    "Finance": [
        ["Finance Analytics", "Financial Data Analyst",
         "Improve profitability, forecasting accuracy, and cost control using financial data",
         "SQL, Excel, Python, Statistics, ML, GenAI, Financial Modeling",
         "JP Morgan, HDFC Bank, American Express, Barclays"],
        ["Finance Analytics", "Risk & Financial Planning Analyst",
         "Predict financial risks, detect anomalies, and strengthen budgeting & planning",
         "SQL, Python, Forecasting, Statistics, ML, GenAI",
         "KPMG, EY, Deloitte, PwC"]
    ],
    "Healthcare": [
        ["Healthcare Analytics", "Healthcare Data Analyst",
         "Analyze patient and hospital data to improve outcomes, efficiency, and care quality",
         "SQL, Python, Statistics, ML, GenAI, Healthcare Data",
         "Apollo Hospitals, Fortis, Practo, Narayana Health"],
        ["Healthcare Analytics", "Clinical Risk & Outcomes Analyst",
         "Predict patient risks, track treatment effectiveness, and optimize resource utilization",
         "SQL, Python, Statistics, ML, GenAI",
         "GE Healthcare, Philips, Medtronic"]
    ],
    "E-Commerce": [
        ["E-Commerce Analytics", "E-Commerce Data Analyst",
         "Optimize sales, pricing, and conversion using customer and product data",
         "SQL, Python, Statistics, ML, GenAI",
         "Amazon, Flipkart, Meesho, Nykaa"],
        ["E-Commerce Analytics", "Customer & Growth Analyst",
         "Analyze customer behavior, churn, and campaign performance to drive growth",
         "SQL, Python, Statistics, ML, GenAI",
         "Myntra, Swiggy, Zomato"]
    ],
    "Supply Chain": [
        ["Supply Chain Analytics", "Supply Chain Data Analyst",
         "Forecast demand and optimize inventory, logistics, and procurement efficiency",
         "SQL, Python, Statistics, ML, GenAI",
         "Amazon, DHL, Flipkart, Delhivery"]
    ],
    "Automobile": [
        ["Automobile Analytics", "Automotive Data Analyst",
         "Analyze vehicle, sensor, and production data to improve quality and efficiency",
         "SQL, Python, Statistics, ML, GenAI, IoT / Telematics Data",
         "Tata Motors, Mahindra, Hyundai, Maruti Suzuki"],
        ["Automobile Analytics", "Manufacturing Operations Analyst",
         "Reduce defects, downtime, and production bottlenecks using analytics",
         "SQL, Python, Time Series, ML, GenAI",
         "Bosch, Continental, TVS Motor, Ashok Leyland"]
    ],
    "Manufacturing": [
        ["Manufacturing Analytics", "Manufacturing Data Analyst",
         "Optimize production output, quality, and operational costs",
         "SQL, Python, Statistics, ML, GenAI, Process Data",
         "Siemens, ABB, GE, Schneider Electric"]
    ],
    "Retail": [
        ["Retail Analytics", "Retail Data Analyst",
         "Optimize inventory, sales forecasting, and customer insights",
         "SQL, Python, Statistics, ML, GenAI",
         "Reliance Retail, DMart, Big Bazaar, Spencer's"]
    ],
    "HR Analytics": [
        ["HR Analytics", "HR Data Analyst",
         "Analyze workforce trends, attrition patterns, and recruitment effectiveness",
         "SQL, Python, Statistics, ML, GenAI",
         "Deloitte, Accenture, IBM, Wipro"]
    ],
    "Cyber Security": [
        ["Cyber Security Analytics", "Security Data Analyst",
         "Detect threats, analyze patterns, and strengthen security posture",
         "SQL, Python, ML, GenAI, SIEM Tools",
         "Cisco, Palo Alto, CrowdStrike, Fortinet"]
    ]
}

def filter_skills_for_technologies(skills_str, technologies):
    """
    Filter the Key Technical Skills cell in the career table based on selected technologies.
    Removes ML, GenAI, etc. if they were not selected. Always keeps SQL, Python, Statistics, Excel.
    """
    if not technologies:
        return skills_str
    tech_lower = set(t.lower() for t in technologies)
    has_ml     = any(t in tech_lower for t in ['machine learning', 'gen ai', 'ml', 'genai'])
    has_bi     = any(t in tech_lower for t in ['power bi', 'tableau', 'visualization'])
    has_cloud  = any(t in tech_lower for t in ['cloud', 'etl', 'data engineering'])

    # Split on comma, strip each token
    tokens = [s.strip() for s in skills_str.split(',')]
    filtered = []
    for t in tokens:
        tl = t.lower()
        # Always keep base skills
        if any(base in tl for base in ['sql', 'python', 'statistics', 'excel', 'forecasting',
                                        'financial modeling', 'financial modelling', 'process data',
                                        'healthcare data', 'siem', 'iot', 'telematics', 'time series']):
            filtered.append(t)
        elif any(ml in tl for ml in ['ml', 'machine learning', 'genai', 'gen ai', 'llm',
                                      'rag', 'embedding', 'fine-tun', 'prompt']):
            if has_ml:
                filtered.append(t)
        elif any(bi in tl for bi in ['power bi', 'tableau', 'visualization', 'dashboard']):
            if has_bi:
                filtered.append(t)
        elif any(cl in tl for cl in ['cloud', 'etl', 'data engineering', 'warehouse']):
            if has_cloud:
                filtered.append(t)
        else:
            filtered.append(t)
    return ', '.join(filtered) if filtered else skills_str



def get_ai_prescription_text(selected_domains, technologies=None):
    if technologies is None:
        technologies = []
    domain_str = " & ".join(selected_domains)
    tech_str = ", ".join(technologies) if technologies else "SQL, Python, Statistics"
    has_ml = any(t.lower() in ['machine learning', 'gen ai', 'ml', 'genai'] for t in technologies)

    if not GROQ_API_KEY or GROQ_API_KEY == "your_groq_api_key_here":
        return {"error": "API Key not configured"}

    ml_instruction = (
        "Include references to Machine Learning and GenAI where relevant."
        if has_ml else
        "Do NOT mention Machine Learning, GenAI, LLMs, RAG, or any AI/ML terms. Focus only on SQL, statistics, and data analysis."
    )

    PROMPT = f"""You are a Senior Data Scientist at Analytics Avenue.
Generate a JSON prescription for a candidate targeting: {domain_str}
Technologies selected: {tech_str}
{ml_instruction}

CRITICAL RULES:
1. Use <b>text</b> for bold formatting on domain names and technologies.
2. Every sentence must be grammatically correct with proper capitalisation.
3. Return ONLY these keys:
   - "intro_line": One sentence introduction with <b> tags
   - "domain_bullets": List of domain descriptions with <b> tags (one bullet per domain)
   - "projects_bullet": Hands-on projects description. Only mention GenAI/ML if those technologies were selected.
   - "final_sentence": Closing sentence listing only the selected technologies with <b> tags

Example for "Finance & Supply Chain" with SQL, Python, Machine Learning, Gen AI:
{{
  "intro_line": "Given your background, we will support your transition into <b>Finance & Supply Chain Analytics</b> roles, enabling you to solve real business and operational problems using <b>Machine Learning and GenAI</b>.",
  "domain_bullets": [
    "In <b>Finance Analytics</b>, you will work on financial performance analysis, budgeting and forecasting, risk assessment, fraud pattern identification, and profitability optimisation.",
    "In <b>Supply Chain Analytics</b>, you will focus on demand forecasting, inventory optimisation, logistics performance, supplier analysis, and end-to-end cost efficiency."
  ],
  "projects_bullet": "Hands-on projects include financial variance and profitability analysis, risk and anomaly detection, demand forecasting, inventory health analysis, logistics optimisation, and supplier performance tracking. You will also leverage <b>GenAI</b> for automated insights, root-cause analysis, and conversational analytics (projects revealed during placement training).",
  "final_sentence": "You will apply <b>SQL, Python, Machine Learning, and GenAI</b> to <b>Finance and Supply Chain</b> datasets, preparing you for high-impact analytics roles across these domains."
}}
NOW GENERATE for: {domain_str} with technologies: {tech_str}
Match the style, grammar, and capitalisation of the example above. Return ONLY valid JSON."""

    try:
        client = Groq(api_key=GROQ_API_KEY)
        completion = client.chat.completions.create(
            messages=[{"role": "user", "content": PROMPT}],
            model="llama-3.3-70b-versatile",
            temperature=0.3,
            response_format={"type": "json_object"}
        )
        data = json.loads(completion.choices[0].message.content)
        data["domains_title"] = domain_str
        return data
    except Exception as e:
        return {"error": str(e)}


def get_table_data_with_rowspan(selected_domains, technologies=None):
    if technologies is None:
        technologies = []
    table_rows = []
    domain_rowspan_map = {}
    for domain in selected_domains:
        if domain in CAREER_TEMPLATES:
            domain_rows = CAREER_TEMPLATES[domain]
            domain_rowspan_map[domain] = len(domain_rows)
            for row in domain_rows:
                # row = [domain_label, role, challenge, skills, companies]
                filtered_skills = filter_skills_for_technologies(row[3], technologies)
                table_rows.append([row[0], row[1], row[2], filtered_skills, row[4]])
    return table_rows, domain_rowspan_map


# ==========================================
# PDF HELPER FUNCTIONS
# ==========================================
def draw_header_no_line(c, page_width, page_height):
    header_path = "assets/header.png"
    if not os.path.exists(header_path):
        c.setFillColor(colors.red)
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, page_height - 50, "ERROR: header.png not found!")
        return 100
    try:
        img = Image.open(header_path)
        img_width, img_height = img.size
        aspect_ratio = img_width / img_height

        # Use full content width for maximum sharpness — no small scaling
        L = MARGINS['left']
        R = MARGINS['right']
        header_width  = page_width - L - R
        header_height = header_width / aspect_ratio

        x_pos = L
        y_pos = page_height - header_height - 10

        c.drawImage(
            header_path, x_pos, y_pos,
            width=header_width, height=header_height,
            preserveAspectRatio=True, mask='auto'
        )

        # Horizontal line below header — same as page 3
        line_y = y_pos - 6
        c.setStrokeColor(colors.black)
        c.setLineWidth(1)
        c.line(L, line_y, page_width - R, line_y)

        return header_height + 30

    except Exception as e:
        return 100


# ==========================================
# PAGE 1 — PDF
# ==========================================
def create_page1(c, name, status, ai_content, program="", technologies=None, consultation_note=""):
    if technologies is None:
        technologies = ["SQL", "Python", "Statistics", "Power BI", "Machine Learning", "Gen AI"]
    if not program:
        program = "Nationwide Data Analytics Training and Placement Program 2026"
    page_width, page_height = A4
    L = MARGINS['left']
    R = page_width - MARGINS['right']
    W = R - L


    header_space = draw_header_no_line(c, page_width, page_height)
    y = page_height - header_space - 10

    style_normal = ParagraphStyle('Normal', fontName='Times-Roman', fontSize=11, leading=15, alignment=TA_LEFT)
    style_bullet_text = ParagraphStyle('BulletText', fontName='Times-Roman', fontSize=11, leading=15, alignment=TA_LEFT)

    def draw_bullet_list(c, items, y, L, W, style):
        """Render bullet list using a 2-col Table: bullet | text. Guarantees alignment."""
        BULLET_COL = 14
        TEXT_COL   = W - BULLET_COL
        data = []
        for item in items:
            data.append([
                Paragraph("•", style),
                Paragraph(item, style)
            ])
        tbl = Table(data, colWidths=[BULLET_COL, TEXT_COL])
        tbl.setStyle(TableStyle([
            ('VALIGN',        (0,0), (-1,-1), 'TOP'),
            ('ALIGN',         (0,0), (0,-1),  'LEFT'),
            ('LEFTPADDING',   (0,0), (-1,-1), 0),
            ('RIGHTPADDING',  (0,0), (-1,-1), 0),
            ('TOPPADDING',    (0,0), (-1,-1), 0),
            ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ]))
        tbl.wrapOn(c, W, 1000)
        th = tbl._height
        tbl.drawOn(c, L, y - th)
        return y - th

    # ── Hi name ──
    c.setFillColor(colors.black)
    c.setFont('Times-Bold', 11)
    c.drawString(L, y, f"Hi {name},")
    y -= 18

    # ── Intro sentence ──
    intro_text = (
        "Our Senior Data Scientist <b>Mr. Subramani</b>, has shared with you the "
        f"prescription based on your recent consultation to join our "
        f"<b>{program}</b>."
    )
    p = Paragraph(intro_text, style_normal)
    _, ph = p.wrap(W, 300)
    p.drawOn(c, L, y - ph)
    y -= ph + 14

    # ── About Us heading ──
    c.setFont('Times-Bold', 11)
    c.drawString(L, y, "About Us")
    y -= 14

    # ── About Us body ──
    about_text = (
        "At <b>Analytics Avenue and Advanced Analytics</b>, we are a team of "
        "<b>Data Scientists, Data Engineers, and BI Developers</b> throughout India "
        "across various MNCs joined together to keep a pause for unemployment and "
        "empowered <b>500+ professionals</b> in the past year, enabling them to "
        "transition into various <b>Data Analytics roles</b>."
    )
    p = Paragraph(about_text, style_normal)
    _, h = p.wrap(W, 200)
    p.drawOn(c, L, y - h)
    y -= h + 12

    # ── Instruction line ──
    instr_text = "Below you can find the key outcomes and suggestions given by our Data Scientist"
    p_instr = Paragraph(f"<b>{instr_text}</b>", style_normal)
    _, h = p_instr.wrap(W, 60)
    p_instr.drawOn(c, L, y - h)
    y -= (h + 6)

    # ── Details block using Table for perfect label/value alignment ──
    details_data = [
        [Paragraph(f"<b>Name</b>",                 style_normal), Paragraph("<b>:</b>", style_normal), Paragraph(name,                                                             style_normal)],
        [Paragraph(f"<b>Status</b>",               style_normal), Paragraph("<b>:</b>", style_normal), Paragraph(status,                                                           style_normal)],
        [Paragraph(f"<b>Technologies Needed</b>",  style_normal), Paragraph("<b>:</b>", style_normal), Paragraph(", ".join(technologies),                                          style_normal)],
        [Paragraph(f"<b>Sectors Covered</b>",      style_normal), Paragraph("<b>:</b>", style_normal), Paragraph(ai_content.get('domains_title', 'Finance & Supply Chain'),        style_normal)],
    ]
    col1_w = 148
    col2_w = 16
    col3_w = W - col1_w - col2_w
    details_table = Table(details_data, colWidths=[col1_w, col2_w, col3_w])
    details_table.setStyle(TableStyle([
        ('VALIGN',        (0, 0), (-1, -1), 'TOP'),
        ('ALIGN',         (0, 0), (-1, -1), 'LEFT'),
        ('LEFTPADDING',   (0, 0), (-1, -1), 0),
        ('RIGHTPADDING',  (0, 0), (-1, -1), 0),
        ('TOPPADDING',    (0, 0), (-1, -1), 2),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ('LINEBELOW',     (0, -1), (-1, -1), 0, colors.white),
    ]))
    details_table.wrapOn(c, W, 200)
    dt_h = details_table._height
    details_table.drawOn(c, L, y - dt_h)
    y -= dt_h + 8

    # ── Key Outcomes heading ──
    c.setFont('Times-Bold', 11)
    c.drawString(L, y, "Key Outcomes")
    y -= 14

    # Dynamic key outcomes based on selected technologies
    tech_set = set(t.lower() for t in technologies)
    outcomes = []
    if any(t in tech_set for t in ['sql', 'python', 'statistics', 'power bi']):
        techs = [t for t in technologies if t.lower() in ['sql', 'python', 'statistics', 'power bi']]
        outcomes.append(f"Data Analysis Skills ({', '.join(techs)}, Visualization)")
    if any(t in tech_set for t in ['cloud', 'etl', 'data engineering']):
        techs = [t for t in technologies if t.lower() in ['cloud', 'etl', 'data engineering', 'sql', 'python']]
        outcomes.append(f"Data Engineer Skills ({', '.join(techs)}, Data Warehousing, ETL Orchestration)")
    if any(t in tech_set for t in ['machine learning', 'gen ai', 'ml', 'genai']):
        outcomes.append("Machine Learning & Gen AI (LLMs, Prompt Engineering, RAG Pipelines, Embeddings, Vector Databases, Fine-Tuning & Deployment)")
    if any(t in tech_set for t in ['power bi', 'tableau', 'visualization']):
        techs = [t for t in technologies if t.lower() in ['power bi', 'tableau', 'visualization']]
        outcomes.append(f"Business Intelligence & Visualization ({', '.join(techs)}, Dashboard Development)")
    outcomes.append(f"Domain Knowledge ({ai_content.get('domains_title', 'Finance & Supply Chain')})")
    outcomes.append("Recreate industrial-standard projects worked on by our Data Scientists")
    outcomes.append("Placement opportunities, organic job calls and referral drives")
    y = draw_bullet_list(c, outcomes, y, L, W, style_normal)

    y -= 10
    # ── Prescription heading ──
    c.setFont('Times-Bold', 11)
    c.drawString(L, y, "Prescription:")
    y -= 14

    # If a custom consultation note was typed, render it first
    if consultation_note and consultation_note.strip():
        note_clean = consultation_note.strip().strip('"').strip("'").strip()
        note_clean = note_clean[0].upper() + note_clean[1:]
        if note_clean and note_clean[-1] not in ".!?":
            note_clean += "."
        p_note = Paragraph(note_clean, style_normal)
        _, nh = p_note.wrap(W, 400)
        p_note.drawOn(c, L, y - nh)
        y -= (nh + 6)

    # ── Intro line ──
    intro_line = ai_content.get('intro_line', "Given your background...")
    p = Paragraph(intro_line, style_normal)
    _, h = p.wrap(W, 400)
    p.drawOn(c, L, y - h)
    y -= (h + 6)

    # ── Domain bullets + projects bullet ──
    presc_bullets = []
    for b_text in ai_content.get('domain_bullets', []):
        if b_text.strip():
            presc_bullets.append(b_text)
    projects_bullet = ai_content.get('projects_bullet', '')
    if projects_bullet:
        presc_bullets.append(projects_bullet)
    if presc_bullets:
        y = draw_bullet_list(c, presc_bullets, y, L, W, style_normal)
        y -= 2

    # ── Final sentence ──
    final_sentence = ai_content.get('final_sentence', "")
    if final_sentence:
        p_final = Paragraph(final_sentence, style_normal)
        _, fh = p_final.wrap(W, 400)
        p_final.drawOn(c, L, y - fh)
        y -= fh




# ==========================================
# PAGE 2 — PDF
# ==========================================
def create_page2(c, ai_content, table_rows, domain_rowspan_map, technologies=None):
    if technologies is None:
        technologies = []
    tech_set_p2 = set(t.lower() for t in technologies)
    has_ml_p2    = any(t in tech_set_p2 for t in ['machine learning', 'gen ai', 'ml', 'genai'])
    has_bi_p2    = any(t in tech_set_p2 for t in ['power bi', 'tableau', 'visualization'])

    # Build dynamic project description for Industry-Relevant Projects
    focus_items = ["data modelling", "EDA"]
    if has_ml_p2:
        focus_items.append("Machine Learning")
        focus_items.append("GenAI")
    if has_bi_p2:
        focus_items.append("Business Intelligence")
    focus_items += ["forecasting", "cost optimisation", "anomaly detection", "decision support"]
    focus_str = ", ".join(focus_items)
    industry_projects_text = (
        f"Work on 3 projects across {ai_content['domains_title']}, focusing on "
        f"{focus_str}."
    )

    page_width, page_height = A4


    header_space = draw_header_no_line(c, page_width, page_height)
    L = MARGINS['left']
    R = page_width - MARGINS['right']
    W = R - L
    y = page_height - header_space - 15

    style_small = ParagraphStyle('Small', fontName='Times-Roman', fontSize=11, leading=13, alignment=TA_LEFT)
    style_heading = ParagraphStyle('Heading', fontName='Times-Bold', fontSize=11, leading=13, alignment=TA_LEFT)

    c.setFillColor(colors.black)
    c.setFont('Times-Bold', 11)
    c.drawString(L, y, "Our Customized Services for you:")
    y -= 14

    services_data = [
        [Paragraph("<b>Service</b>", style_heading), Paragraph("<b>Details</b>", style_heading)],
        [Paragraph("<b>1. Industry-Relevant Projects</b>", style_heading),
         Paragraph(industry_projects_text, style_small)],
        [Paragraph("<b>2. Secret Job Portals Access</b>", style_heading),
         Paragraph("Setup and optimize your profile on 9 exclusive job portals to help you receive organic job calls", style_small)],
        [Paragraph("<b>3. Interview Preparation Materials</b>", style_heading),
         Paragraph("Lifetime access to interview notes, preparation guides, and materials prepared by top Data Scientists in real interview scenarios", style_small)],
        [Paragraph("<b>4. Monthly In-Person Training</b>", style_heading),
         Paragraph("Attend monthly in-house classroom sessions (1 weekend per month) for revision, rapid preparation, and mentorship from experienced professionals", style_small)]
    ]

    services_table = Table(services_data, colWidths=[W * 0.32, W * 0.68])
    services_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.white),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 5),
        ('TOPPADDING', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
    ]))
    services_table.wrapOn(c, W, page_height)
    services_h = services_table._height
    services_table.drawOn(c, L, y - services_h)
    y -= (services_h + 12)

    c.setFont("Times-Bold", 11)
    c.drawString(L, y, f"{ai_content['domains_title']} \u2013 Career Prescription Table")
    y -= 14
    c.setFont("Times-Italic", 11)
    c.drawString(L, y, "(Actual projects will be revealed during placement training)")
    y -= 14

    headers = ["Domain", "Role", "Exciting Challenge", "Key Technical Skills", "Targeted Companies"]
    career_data = [[Paragraph(f"<b>{h}</b>", style_heading) for h in headers]]

    current_row = 1
    processed_domains = {}
    full_domain_rowspan = {}

    for row in table_rows:
        full_domain = row[0]
        if full_domain not in processed_domains:
            processed_domains[full_domain] = current_row
            full_domain_rowspan[full_domain] = 1
            career_data.append([
                Paragraph(f"<b>{full_domain}</b>", style_small),
                Paragraph(str(row[1]), style_small),
                Paragraph(str(row[2]), style_small),
                Paragraph(str(row[3]), style_small),
                Paragraph(str(row[4]), style_small)
            ])
        else:
            full_domain_rowspan[full_domain] += 1
            career_data.append([
                "",
                Paragraph(str(row[1]), style_small),
                Paragraph(str(row[2]), style_small),
                Paragraph(str(row[3]), style_small),
                Paragraph(str(row[4]), style_small)
            ])
        current_row += 1

    col_widths = [W * 0.18, W * 0.15, W * 0.25, W * 0.20, W * 0.22]
    career_table = Table(career_data, colWidths=col_widths, repeatRows=1)

    table_style = [
        ('BACKGROUND', (0, 0), (-1, 0), colors.white),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('BOX', (0, 0), (-1, -1), 1, colors.black),
        ('LINEBELOW', (0, 0), (-1, 0), 1, colors.black),
        ('INNERGRID', (1, 0), (-1, -1), 0.5, colors.black),
        ('LINEAFTER', (0, 0), (0, -1), 1, colors.black),
        ('LEFTPADDING', (0, 0), (-1, -1), 5),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
    ]

    for full_domain, start_row in processed_domains.items():
        rowspan = full_domain_rowspan[full_domain]
        if rowspan > 1:
            table_style.append(('SPAN', (0, start_row), (0, start_row + rowspan - 1)))
            table_style.append(('VALIGN', (0, start_row), (0, start_row + rowspan - 1), 'MIDDLE'))
            for sub_row in range(start_row, start_row + rowspan - 1):
                table_style.append(('LINEBELOW', (1, sub_row), (-1, sub_row), 0.3, colors.lightgrey))
        table_style.append(('LINEBELOW', (0, start_row + rowspan - 1), (-1, start_row + rowspan - 1), 1.5, colors.black))

    career_table.setStyle(TableStyle(table_style))
    career_table.wrapOn(c, W, page_height)
    career_h = career_table._height
    career_table.drawOn(c, L, y - career_h)


# ==========================================
# PDF GENERATION
# ==========================================
def create_final_pdf(name, status, ai_content, table_rows, domain_rowspan_map, output_path, program="", technologies=None, consultation_note=""):
    if technologies is None:
        technologies = ["SQL", "Python", "Statistics", "Power BI", "Machine Learning", "Gen AI"]
    if not program:
        program = "Nationwide Data Analytics Training and Placement Program 2026"
    buffer1 = io.BytesIO()
    c1 = canvas.Canvas(buffer1, pagesize=A4)
    create_page1(c1, name, status, ai_content, program, technologies, consultation_note)
    c1.save()
    buffer1.seek(0)

    buffer2 = io.BytesIO()
    c2 = canvas.Canvas(buffer2, pagesize=A4)
    create_page2(c2, ai_content, table_rows, domain_rowspan_map, technologies)
    c2.save()
    buffer2.seek(0)

    writer = PdfWriter()
    reader1 = PdfReader(buffer1)
    writer.add_page(reader1.pages[0])
    reader2 = PdfReader(buffer2)
    writer.add_page(reader2.pages[0])

    template_path = "assets/template.pdf"
    if os.path.exists(template_path):
        template_reader = PdfReader(template_path)
        if len(template_reader.pages) >= 3:
            page3_template = template_reader.pages[2]
            if page3_template.mediabox.width != A4[0] or page3_template.mediabox.height != A4[1]:
                page3_template.scale_to(A4[0], A4[1])
            writer.add_page(page3_template)

    with open(output_path, 'wb') as f:
        writer.write(f)

    return True, None


# ==========================================
# DOCX HELPER — set cell border
# ==========================================
def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_bold_run(para, text, size_pt=11, color_hex=None):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size_pt)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return run


def add_run(para, text, bold=False, size_pt=11, italic=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    return run


def parse_bold_text(para, html_text, size_pt=11):
    """Parse <b>...</b> tags and add runs with correct bold formatting."""
    import re
    parts = re.split(r'(<b>.*?</b>)', html_text)
    for part in parts:
        if part.startswith('<b>') and part.endswith('</b>'):
            inner = part[3:-4]
            r = para.add_run(inner)
            r.bold = True
            r.font.size = Pt(size_pt)
        else:
            clean = part.replace('</b>', '').replace('<b>', '')
            if clean:
                r = para.add_run(clean)
                r.bold = False
                r.font.size = Pt(size_pt)


# ==========================================
# WORD DOCUMENT GENERATION
# ==========================================
def create_word_doc(name, status, ai_content, table_rows, domain_rowspan_map, output_path, program="", technologies=None, consultation_note=""):
    if technologies is None:
        technologies = ["SQL", "Python", "Statistics", "Power BI", "Machine Learning", "Gen AI"]
    if not program:
        program = "Nationwide Data Analytics Training and Placement Program 2026"
    doc = Document()

    # ── Page setup: A4, margins matching PDF ──
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)
    section.top_margin    = Cm(1.2)
    section.bottom_margin = Cm(1.8)

    # ── Default style ──
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    # ── HEADER IMAGE ──
    header_path = "assets/header.png"
    if os.path.exists(header_path):
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.add_run()
        run.add_picture(header_path, width=Inches(6.5))
        header_para.paragraph_format.space_after = Pt(6)

    # ── Divider line ──
    div_para = doc.add_paragraph()
    div_para.paragraph_format.space_before = Pt(0)
    div_para.paragraph_format.space_after = Pt(8)
    pPr = div_para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom_bdr = OxmlElement('w:bottom')
    bottom_bdr.set(qn('w:val'), 'single')
    bottom_bdr.set(qn('w:sz'), '6')
    bottom_bdr.set(qn('w:color'), '000000')
    pBdr.append(bottom_bdr)
    pPr.append(pBdr)

    # ── Hi name ──
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_bold_run(p, f"Hi {name},")

    # ── Intro paragraph ──
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    parse_bold_text(p,
        f"Our Senior Data Scientist <b>Mr. Subramani</b>, has shared with you the prescription "
        f"based on your recent consultation to join our "
        f"<b>{program}</b>."
    )

    # ── About Us ──
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    add_bold_run(p, "About Us")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    parse_bold_text(p,
        "At <b>Analytics Avenue and Advanced Analytics</b>, we are a team of "
        "<b>Data Scientists, Data Engineers, and BI Developers</b> throughout India "
        "across various MNCs joined together to keep a pause for unemployment and "
        "empowered <b>500+ professionals</b> in the past year, enabling them to "
        "transition into various <b>Data Analytics roles</b>."
    )

    # ── Instruction line ──
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    add_bold_run(p, "Below you can find the key outcomes and suggestions given by our Data Scientist")

    # ── Details table ──
    details = [
        ("Name", name),
        ("Status", status),
        ("Technologies Needed", ", ".join(technologies)),
        ("Sectors Covered", ai_content.get('domains_title', 'Finance & Supply Chain'))
    ]
    det_table = doc.add_table(rows=len(details), cols=3)
    det_table.style = 'Table Grid'
    det_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    col_w = [Cm(4.5), Cm(0.5), Cm(11)]
    for i, (label, value) in enumerate(details):
        row = det_table.rows[i]
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(0.5)
        row.cells[2].width = Cm(11)
        p0 = row.cells[0].paragraphs[0]
        add_bold_run(p0, label)
        p1 = row.cells[1].paragraphs[0]
        add_bold_run(p1, ":")
        p2 = row.cells[2].paragraphs[0]
        add_run(p2, value)
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            # remove borders for clean look matching PDF
            set_cell_border(cell,
                top={'val': 'none'}, bottom={'val': 'none'},
                left={'val': 'none'}, right={'val': 'none'}
            )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    # ── Key Outcomes ──
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    add_bold_run(p, "Key Outcomes")

    # Dynamic outcomes based on selected technologies
    tech_set_doc = set(t.lower() for t in technologies)
    outcomes = []
    if any(t in tech_set_doc for t in ['sql', 'python', 'statistics', 'power bi']):
        techs = [t for t in technologies if t.lower() in ['sql', 'python', 'statistics', 'power bi']]
        outcomes.append(f"Data Analysis Skills ({', '.join(techs)}, Visualization)")
    if any(t in tech_set_doc for t in ['cloud', 'etl', 'data engineering']):
        techs = [t for t in technologies if t.lower() in ['cloud', 'etl', 'data engineering', 'sql', 'python']]
        outcomes.append(f"Data Engineer Skills ({', '.join(techs)}, Data Warehousing, ETL orchestration)")
    if any(t in tech_set_doc for t in ['machine learning', 'gen ai', 'ml', 'genai']):
        outcomes.append("Machine Learning & Gen AI (LLMs, Prompt Engineering, RAG Pipelines, Embeddings, Vector Databases, Fine-Tuning & Deployment)")
    if any(t in tech_set_doc for t in ['power bi', 'tableau', 'visualization']):
        techs = [t for t in technologies if t.lower() in ['power bi', 'tableau', 'visualization']]
        outcomes.append(f"Business Intelligence & Visualization ({', '.join(techs)}, Dashboard Development)")
    outcomes.append(f"Domain Knowledge ({ai_content.get('domains_title', 'Finance & Supply Chain')})")
    outcomes.append("Recreate industrial-standard projects worked on by our Data Scientists")
    outcomes.append("Placement opportunities, Organic job calls and referral drives")
    for item in outcomes:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        add_run(p, item)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    # ── Prescription ──
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(5)
    add_bold_run(p, "Prescription:")

    # Custom consultation note — rendered first if present
    if consultation_note and consultation_note.strip():
        note_clean = consultation_note.strip().strip('"').strip("'").strip()
        note_clean = note_clean[0].upper() + note_clean[1:]
        if note_clean and note_clean[-1] not in ".!?":
            note_clean += "."
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        add_run(p, note_clean)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    parse_bold_text(p, ai_content.get('intro_line', ''))

    for b_text in ai_content.get('domain_bullets', []):
        if b_text.strip():
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            parse_bold_text(p, b_text)

    projects = ai_content.get('projects_bullet', '')
    if projects:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        parse_bold_text(p, projects)

    final = ai_content.get('final_sentence', '')
    if final:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        parse_bold_text(p, final)

    # ── PAGE BREAK ──
    doc.add_page_break()

    # ── Header image page 2 ──
    if os.path.exists(header_path):
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.add_run()
        run.add_picture(header_path, width=Inches(6.5))
        header_para.paragraph_format.space_after = Pt(6)

    div_para2 = doc.add_paragraph()
    div_para2.paragraph_format.space_before = Pt(0)
    div_para2.paragraph_format.space_after = Pt(8)
    pPr2 = div_para2._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    b2 = OxmlElement('w:bottom')
    b2.set(qn('w:val'), 'single')
    b2.set(qn('w:sz'), '6')
    b2.set(qn('w:color'), '000000')
    pBdr2.append(b2)
    pPr2.append(pBdr2)

    # ── Customized Services ──
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_bold_run(p, "Our Customized Services for you:")
    # Dynamic Industry-Relevant Projects description
    tech_set_doc2 = set(t.lower() for t in technologies)
    has_ml_doc  = any(t in tech_set_doc2 for t in ['machine learning', 'gen ai', 'ml', 'genai'])
    has_bi_doc  = any(t in tech_set_doc2 for t in ['power bi', 'tableau', 'visualization'])
    focus_doc = ["data modelling", "EDA"]
    if has_ml_doc:
        focus_doc.append("Machine Learning")
        focus_doc.append("GenAI")
    if has_bi_doc:
        focus_doc.append("Business Intelligence")
    focus_doc += ["forecasting", "cost optimisation", "anomaly detection", "decision support"]
    industry_proj_doc = (
        f"Work on 3 projects across {ai_content['domains_title']}, focusing on "
        f"{', '.join(focus_doc)}."
    )

    svc_rows = [
        ("1. Industry-Relevant Projects", industry_proj_doc),
        ("2. Secret Job Portals Access",
         "Setup and optimize your profile on 9 exclusive job portals to help you receive organic job calls"),
        ("3. Interview Preparation Materials",
         "Lifetime access to interview notes, preparation guides, and materials prepared by top Data Scientists in real interview scenarios"),
        ("4. Monthly In-Person Training",
         "Attend monthly in-house classroom sessions (1 weekend per month) for revision, rapid preparation, and mentorship from experienced professionals"),
    ]

    svc_table = doc.add_table(rows=1 + len(svc_rows), cols=2)
    svc_table.style = 'Table Grid'
    svc_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hrow = svc_table.rows[0]
    hrow.cells[0].width = Cm(5)
    hrow.cells[1].width = Cm(11)
    h0 = hrow.cells[0].paragraphs[0]
    add_bold_run(h0, "Service")
    h1 = hrow.cells[1].paragraphs[0]
    add_bold_run(h1, "Details")

    for i, (svc, detail) in enumerate(svc_rows):
        row = svc_table.rows[i + 1]
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(11)
        p0 = row.cells[0].paragraphs[0]
        add_bold_run(p0, svc)
        p1 = row.cells[1].paragraphs[0]
        add_run(p1, detail)
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(4)
            cell.paragraphs[0].paragraph_format.space_after = Pt(4)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)

    # ── Career Prescription Table title ──
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_bold_run(p, f"{ai_content['domains_title']} – Career Prescription Table")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("(Actual projects will be revealed during placement training)")
    r.italic = True
    r.font.size = Pt(10)

    # ── Career Table ──
    career_headers = ["Domain", "Role", "Exciting Challenge", "Key Technical Skills", "Targeted Companies"]
    total_rows = 1 + len(table_rows)
    ct = doc.add_table(rows=total_rows, cols=5)
    ct.style = 'Table Grid'
    ct.alignment = WD_TABLE_ALIGNMENT.LEFT

    col_widths_cm = [3.0, 2.5, 4.2, 3.4, 3.6]

    # Header
    hr = ct.rows[0]
    for j, hdr in enumerate(career_headers):
        cell = hr.cells[j]
        cell.width = Cm(col_widths_cm[j])
        ph = cell.paragraphs[0]
        add_bold_run(ph, hdr, size_pt=10)
        cell.paragraphs[0].paragraph_format.space_before = Pt(3)
        cell.paragraphs[0].paragraph_format.space_after = Pt(3)

    # Data rows
    processed_domains_docx = {}
    full_domain_rowspan_docx = {}
    for row_data in table_rows:
        fd = row_data[0]
        if fd not in full_domain_rowspan_docx:
            full_domain_rowspan_docx[fd] = 1
        else:
            full_domain_rowspan_docx[fd] += 1

    current_row_idx = 1
    domain_start_idx = {}
    for row_data in table_rows:
        fd = row_data[0]
        row = ct.rows[current_row_idx]
        for j in range(5):
            row.cells[j].width = Cm(col_widths_cm[j])

        # Domain cell
        if fd not in domain_start_idx:
            domain_start_idx[fd] = current_row_idx
            pd = row.cells[0].paragraphs[0]
            add_bold_run(pd, fd, size_pt=10)

        for j, val in enumerate(row_data[1:], start=1):
            pc = row.cells[j].paragraphs[0]
            add_run(pc, str(val), size_pt=10)

        for j in range(5):
            row.cells[j].paragraphs[0].paragraph_format.space_before = Pt(3)
            row.cells[j].paragraphs[0].paragraph_format.space_after = Pt(3)

        current_row_idx += 1

    # Merge domain cells vertically
    for fd, start_idx in domain_start_idx.items():
        span = full_domain_rowspan_docx[fd]
        if span > 1:
            end_idx = start_idx + span - 1
            ct.cell(start_idx, 0).merge(ct.cell(end_idx, 0))
            merged_cell = ct.cell(start_idx, 0)
            merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ── PAGE 3 from template — new section (section break acts as page break) ──
    template_path_docx = "assets/template.pdf"
    if os.path.exists(template_path_docx):
        try:
            from pdf2image import convert_from_bytes
            with open(template_path_docx, "rb") as tf:
                pdf_bytes = tf.read()
            pages = convert_from_bytes(pdf_bytes, first_page=3, last_page=3, dpi=130)
            if pages:
                img_buf = io.BytesIO()
                pages[0].save(img_buf, format="PNG")
                img_buf.seek(0)
                # New section with zero margins so image fills full A4 page
                sec3 = doc.add_section()
                sec3.page_width    = Cm(21)
                sec3.page_height   = Cm(29.7)
                sec3.left_margin   = Cm(0)
                sec3.right_margin  = Cm(0)
                sec3.top_margin    = Cm(0)
                sec3.bottom_margin = Cm(0)
                p3_para = doc.add_paragraph()
                p3_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p3_para.paragraph_format.space_before = Pt(0)
                p3_para.paragraph_format.space_after  = Pt(0)
                run_img = p3_para.add_run()
                run_img.add_picture(img_buf, width=Cm(21), height=Cm(29.7))
        except Exception as e:
            doc.add_page_break()
            p3_para = doc.add_paragraph()
            p3_para.add_run(f"[Page 3 could not be embedded: {e}]")



    doc.save(output_path)
    return True, None


# ==========================================
# SEND MAIL FUNCTION
# ==========================================
def send_mail_with_pdf(to_email, cc_emails, subject, body, pdf_path, candidate_name):
    if not GMAIL_USER or not GMAIL_PASSWORD:
        return False, "Gmail credentials not configured in Streamlit secrets (GMAIL_USER, GMAIL_PASSWORD)"

    try:
        msg = MIMEMultipart()
        msg['From'] = GMAIL_USER
        msg['To'] = to_email
        if cc_emails:
            msg['Cc'] = ", ".join(cc_emails)
        msg['Subject'] = subject

        msg.attach(MIMEText(body, 'plain'))

        with open(pdf_path, 'rb') as f:
            pdf_attachment = MIMEApplication(f.read(), _subtype='pdf')
            pdf_attachment.add_header(
                'Content-Disposition', 'attachment',
                filename=os.path.basename(pdf_path)
            )
            msg.attach(pdf_attachment)

        all_recipients = [to_email] + (cc_emails if cc_emails else [])

        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
            server.login(GMAIL_USER, GMAIL_PASSWORD)
            server.sendmail(GMAIL_USER, all_recipients, msg.as_string())

        return True, None
    except Exception as e:
        return False, str(e)


# ==========================================
# STREAMLIT UI
# ==========================================
st.set_page_config(page_title="Analytics Avenue Generator", layout="wide")

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700;800;900&display=swap');
    html, body, [class*="css"] { font-family: 'Inter', sans-serif !important; }
    .block-container { padding-top: 2rem !important; padding-left: 3rem !important; padding-right: 3rem !important; max-width: 100% !important; }
    #MainMenu, footer, header { visibility: hidden; }
    .brand-wrap { display: flex; align-items: center; gap: 18px; margin-bottom: 28px; }
    .brand-name { font-size: 26px; font-weight: 800; color: #064b86; line-height: 1.3; }
    .divider { border: none; border-top: 2px solid #e0e0e0; margin: 0 0 32px 0; }
    h1 { font-size: 48px !important; font-weight: 900 !important; color: #0a0a0a !important; letter-spacing: -1px !important; line-height: 1.1 !important; margin-bottom: 6px !important; }
    .subtitle { font-size: 17px; font-weight: 500; color: #555; margin-bottom: 36px; }
    h2 { font-size: 30px !important; font-weight: 800 !important; color: #0a0a0a !important; margin-bottom: 16px !important; }
    h3 { font-size: 22px !important; font-weight: 700 !important; color: #0a0a0a !important; margin-bottom: 12px !important; }
    .card { background: #fff; border: 1.5px solid #e5e7eb; border-radius: 10px; padding: 24px 28px; margin-bottom: 20px; }
    .card-label { font-size: 13px; font-weight: 700; color: #064b86; text-transform: uppercase; letter-spacing: 1px; margin-bottom: 8px; }
    .card-text { font-size: 16px; font-weight: 500; color: #222; line-height: 1.7; }
    .card ul { margin: 0; padding-left: 18px; }
    .card ul li { font-size: 15px; font-weight: 500; color: #333; margin-bottom: 6px; line-height: 1.6; }
    .stTabs [data-baseweb="tab-list"] { gap: 0px; border-bottom: 2px solid #e0e0e0; margin-bottom: 32px; }
    .stTabs [data-baseweb="tab"] { font-size: 16px !important; font-weight: 600 !important; color: #555 !important; padding: 12px 28px !important; border: none !important; background: transparent !important; }
    .stTabs [aria-selected="true"] { color: #064b86 !important; font-weight: 800 !important; border-bottom: 3px solid #064b86 !important; }
    .stTextInput label, .stSelectbox label, .stMultiSelect label { font-size: 15px !important; font-weight: 700 !important; color: #0a0a0a !important; letter-spacing: 0.3px !important; }
    .stTextInput input { font-size: 16px !important; font-weight: 500 !important; padding: 12px 14px !important; border: 1.5px solid #d0d7de !important; border-radius: 6px !important; background: #fff !important; }
    .stFormSubmitButton > button { background-color: #064b86 !important; color: #fff !important; font-size: 17px !important; font-weight: 700 !important; padding: 14px 36px !important; border-radius: 6px !important; border: none !important; letter-spacing: 0.3px !important; margin-top: 12px !important; width: auto !important; }
    .stFormSubmitButton > button:hover { background-color: #053d70 !important; }
    .stDownloadButton > button { background-color: #1a7f37 !important; color: #fff !important; font-size: 16px !important; font-weight: 700 !important; padding: 12px 28px !important; border-radius: 6px !important; border: none !important; }
    .stAlert p { font-size: 15px !important; font-weight: 600 !important; }
    .streamlit-expanderHeader p { font-size: 17px !important; font-weight: 700 !important; color: #0a0a0a !important; }
    .mail-box { background: #f0f6ff; border: 1.5px solid #bed0f7; border-radius: 10px; padding: 22px 26px; margin-top: 20px; }
</style>
""", unsafe_allow_html=True)

# ── BRAND HEADER ──
logo_url = "https://raw.githubusercontent.com/Analytics-Avenue/streamlit-dataapp/main/logo.png"
st.markdown(f"""
<div class="brand-wrap">
    <img src="{logo_url}" width="64" style="border-radius:8px;">
    <div class="brand-name">Analytics Avenue &amp;<br>Advanced Analytics</div>
</div>
<hr class="divider">
""", unsafe_allow_html=True)

st.title("🤖 AI Prescription Generator")
st.markdown('<p class="subtitle">Generate a personalised data career prescription powered by AI</p>', unsafe_allow_html=True)

tab1, tab2, tab3 = st.tabs(["Overview", "Application", "Special Consultation"])

# ════════════════════════════════════════════════════════
# TAB 1 — OVERVIEW
# ════════════════════════════════════════════════════════
with tab1:
    st.header("Overview")
    st.markdown("""
    <div class="card">
        <div class="card-label">Purpose</div>
        <div class="card-text">
            Generate personalised, AI-powered career prescriptions for aspiring data professionals —
            combining Groq LLaMA 3.3 70B intelligence with domain-specific career templates to produce
            a structured 3-page PDF and editable Word document covering skills, projects, roles, and targeted companies.
        </div>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns(2, gap="large")
    with col1:
        st.subheader("Capabilities")
        st.markdown("""
        <div class="card"><ul>
            <li>Supports 9 domains — Finance, Healthcare, Supply Chain, E-Commerce, HR Analytics, Automobile, Manufacturing, Retail, and Cyber Security.</li>
            <li>AI generates personalised prescription text with domain-specific bullets using Groq LLaMA 3.3 70B.</li>
            <li>Download as PDF (3 pages) or editable Word (.docx) document.</li>
            <li>Send the prescription directly to the candidate's email with CC support.</li>
            <li>Career table includes roles, challenges, key skills, and targeted companies per domain.</li>
        </ul></div>
        """, unsafe_allow_html=True)
    with col2:
        st.subheader("Business Impact")
        st.markdown("""
        <div class="card"><ul>
            <li>Replace manual prescription writing — generate tailored career documents in seconds.</li>
            <li>Deliver consistent, professional-grade prescriptions to every prospective student.</li>
            <li>Send directly to candidates via email without leaving the app.</li>
            <li>Word export lets consultants make last-minute edits before sharing.</li>
            <li>Scale across hundreds of consultations without additional effort.</li>
        </ul></div>
        """, unsafe_allow_html=True)

# ════════════════════════════════════════════════════════
# TAB 2 — APPLICATION
# ════════════════════════════════════════════════════════
with tab2:
    header_ok = os.path.exists("assets/header.png")
    template_ok = os.path.exists("assets/template.pdf")
    if not (header_ok and template_ok):
        st.error("❌ Missing required assets!")
        st.write(f"{'✅' if header_ok else '❌'} header.png")
        st.write(f"{'✅' if template_ok else '❌'} template.pdf")
        st.stop()

    # ── Initialise all session_state keys once ──
    for _k, _v in {
        "generated":    False,
        "pdf_path":     "",
        "docx_path":    "",
        "base_name":    "",
        "pdf_ok":       False,
        "docx_ok":      False,
        "pdf_err":      "",
        "docx_err":     "",
        "ai_content":   {},
        "table_rows":   [],
        "domain_map":   {},
        "cand_name":    "",
        "program":      "",
        "technologies": [],
        "mail_to":      "",
        "mail_cc":      "",
        "mail_subject": "",
        "mail_body":    "",
        "mail_status":       "",   # "" | "sending" | "sent" | "error"
        "mail_msg":          "",
        "consultation_note": "",
    }.items():
        if _k not in st.session_state:
            st.session_state[_k] = _v

    # ════════════════════════════════
    # GENERATE FORM
    # ════════════════════════════════
    st.subheader("Your Details")

    with st.form("form"):
        program = st.selectbox(
            "Program Name *",
            [
                "Nationwide Data Analytics Training and Placement Program 2026",
                "Nationwide Data Engineering and Gen AI Program 2026",
                "Nationwide Data Analytics and Business Intelligence Program 2026"
            ]
        )
        col1, col2 = st.columns(2, gap="large")
        with col1:
            name = st.text_input("Name *", placeholder="e.g. Student Name")
        with col2:
            status = st.selectbox("Status *", ["Working Professional", "Student", "Job Seeker"])
        technologies = st.multiselect(
            "Technologies Needed *",
            ["SQL", "Python", "Statistics", "Power BI", "Machine Learning", "Gen AI",
             "Cloud", "ETL", "Data Engineering", "Tableau", "Visualization"],
            default=[],
            help="Select technologies for this program"
        )
        domains = st.multiselect(
            "Target Domains *",
            ["Finance", "Supply Chain", "Healthcare", "HR Analytics", "E-Commerce",
             "Automobile", "Manufacturing", "Retail", "Cyber Security"],
            help="Select 1–3 domains"
        )
        submit = st.form_submit_button("🚀 Generate Prescription")

    # ── On Generate click — do all work and save to session_state ──
    if submit:
        errors = []
        if not name:
            errors.append("❌ Name is required")
        if not domains:
            errors.append("❌ Select at least one domain")

        if errors:
            for e in errors:
                st.error(e)
        else:
            with st.spinner("🤖 AI generating prescription..."):
                ai_content = get_ai_prescription_text(domains, technologies)

            if "error" in ai_content:
                st.error(f"AI Error: {ai_content['error']}")
            else:
                with st.spinner("📊 Building career table..."):
                    table_rows, domain_rowspan_map = get_table_data_with_rowspan(domains, technologies)

                with st.spinner("📄 Creating PDF & Word document..."):
                    ts        = int(time.time())
                    safe_name = "".join([c for c in name if c.isalnum() or c in (' ', '_')]).rstrip()
                    base_name = f"Prescription_{safe_name.replace(' ', '_')}_{ts}"
                    os.makedirs("output", exist_ok=True)
                    pdf_path  = f"output/{base_name}.pdf"
                    docx_path = f"output/{base_name}.docx"
                    pdf_ok,  pdf_err  = create_final_pdf(name, status, ai_content, table_rows, domain_rowspan_map, pdf_path, program, technologies, st.session_state.get("consultation_note", ""))
                    docx_ok, docx_err = create_word_doc(name, status, ai_content, table_rows, domain_rowspan_map, docx_path, program, technologies, st.session_state.get("consultation_note", ""))

                # Build default mail body
                default_body = (
                    f"Dear {name},\n\n"
                    f"Thank you for your recent consultation with Analytics Avenue & Advanced Analytics.\n\n"
                    f"As discussed, please find attached your personalised Career Prescription prepared by "
                    f"our Senior Data Scientist Mr. Subramani. This document outlines your tailored roadmap, "
                    f"key outcomes, and domain-specific career opportunities in "
                    f"{ai_content.get('domains_title', 'Data Analytics')}.\n\n"
                    f"Your prescription covers:\n"
                    f"  \u2022 Customised career roadmap across {ai_content.get('domains_title', '')}\n"
                    f"  \u2022 Key technical skills: SQL, Python, Statistics, Power BI, Machine Learning, Gen AI\n"
                    f"  \u2022 Industry-relevant projects and placement support\n\n"
                    f"To take the next step, please register and pay the initial \u20b95,000 to block your seat:\n"
                    f"Payment Link: https://pages.razorpay.com/OpenAnalyticsAvenue\n"
                    f"UPI: aard@uco\n\n"
                    f"Feel free to reach out for any queries.\n\n"
                    f"Warm regards,\n"
                    f"Data Consultant\n"
                    f"Analytics Avenue & Advanced Analytics\n"
                    f"Ph / WhatsApp: 9677298268\n"
                    f"Email: supportteam@analyticsavenue.in"
                )

                # Save everything to session_state
                st.session_state["generated"]    = True
                st.session_state["pdf_path"]     = pdf_path
                st.session_state["docx_path"]    = docx_path
                st.session_state["base_name"]    = base_name
                st.session_state["pdf_ok"]       = pdf_ok
                st.session_state["docx_ok"]      = docx_ok
                st.session_state["pdf_err"]      = pdf_err or ""
                st.session_state["docx_err"]     = docx_err or ""
                st.session_state["ai_content"]   = ai_content
                st.session_state["table_rows"]   = table_rows
                st.session_state["domain_map"]   = domain_rowspan_map
                st.session_state["cand_name"]    = name
                st.session_state["program"]      = program
                st.session_state["technologies"] = technologies
                # Reset mail fields for new prescription
                st.session_state["mail_to"]      = ""
                st.session_state["mail_cc"]      = ""
                st.session_state["mail_subject"] = "Your Career Prescription \u2013 Analytics Avenue & Advanced Analytics"
                st.session_state["mail_body"]    = default_body
                st.session_state["mail_status"]  = ""
                st.session_state["mail_msg"]     = ""

    # ════════════════════════════════
    # RESULTS SECTION
    # always rendered from session_state — survives ANY button click
    # ════════════════════════════════
    if st.session_state["generated"]:
        _pdf_ok   = st.session_state["pdf_ok"]
        _docx_ok  = st.session_state["docx_ok"]
        _pdf_path = st.session_state["pdf_path"]
        _docx_path= st.session_state["docx_path"]
        _base     = st.session_state["base_name"]
        _ai       = st.session_state["ai_content"]
        _rows     = st.session_state["table_rows"]
        _dmap     = st.session_state["domain_map"]
        _cname    = st.session_state["cand_name"]

        st.success("✅ Prescription Generated Successfully!")

        # ── Download buttons ──
        dl_col1, dl_col2, _ = st.columns([2, 2, 3])
        with dl_col1:
            if _pdf_ok and os.path.exists(_pdf_path):
                with open(_pdf_path, "rb") as f:
                    st.download_button(
                        "⬇️ Download PDF", f,
                        file_name=f"{_base}.pdf",
                        mime="application/pdf",
                        key="dl_pdf"
                    )
            else:
                st.error(f"PDF Error: {st.session_state['pdf_err']}")

        with dl_col2:
            if _docx_ok and os.path.exists(_docx_path):
                with open(_docx_path, "rb") as f:
                    st.download_button(
                        "📝 Download Word (.docx)", f,
                        file_name=f"{_base}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="dl_docx"
                    )
            else:
                st.error(f"Word Error: {st.session_state['docx_err']}")

        # ════════════════════════════════
        # SEND MAIL SECTION
        # ════════════════════════════════
        st.markdown("---")
        st.subheader("📧 Send Prescription by Email")

        st.markdown('<div class="mail-box">', unsafe_allow_html=True)

        mc1, mc2 = st.columns(2, gap="large")
        with mc1:
            st.text_input("To *", placeholder="candidate@email.com", key="mail_to")
        with mc2:
            st.text_input(
                "CC (comma-separated)",
                placeholder="cc1@email.com, cc2@email.com",
                key="mail_cc"
            )

        st.text_input("Subject", key="mail_subject")
        st.text_area("Email Body", height=280, key="mail_body")

        st.markdown('</div>', unsafe_allow_html=True)

        # ── Status banner (persists across reruns) ──
        if st.session_state["mail_status"] == "sent":
            st.success(f"✅ {st.session_state['mail_msg']}")
        elif st.session_state["mail_status"] == "error":
            st.error(f"❌ {st.session_state['mail_msg']}")
            st.info("💡 Use a Gmail **App Password** — not your regular Gmail password. "
                    "Generate one at: myaccount.google.com/apppasswords  |  "
                    "Then add GMAIL_USER and GMAIL_PASSWORD to Streamlit secrets.")

        # ── Send button — outside any form so it doesn't clear fields ──
        if st.button("📤 Send Mail", type="primary", key="send_mail_btn"):
            _to = st.session_state["mail_to"].strip()
            if not _to:
                st.session_state["mail_status"] = "error"
                st.session_state["mail_msg"]    = "Please enter a To email address."
            else:
                _cc_raw  = st.session_state["mail_cc"]
                _cc_list = [e.strip() for e in _cc_raw.replace(',', '\n').split('\n') if e.strip()] if _cc_raw.strip() else []

                st.session_state["mail_status"] = "sending"
                with st.spinner("📨 Sending email..."):
                    mail_ok, mail_err = send_mail_with_pdf(
                        to_email    = _to,
                        cc_emails   = _cc_list,
                        subject     = st.session_state["mail_subject"],
                        body        = st.session_state["mail_body"],
                        pdf_path    = _pdf_path,
                        candidate_name = _cname
                    )

                if mail_ok:
                    disp = _to + (f"  |  CC: {', '.join(_cc_list)}" if _cc_list else "")
                    st.session_state["mail_status"] = "sent"
                    st.session_state["mail_msg"]    = f"Email sent successfully to {disp}"
                else:
                    st.session_state["mail_status"] = "error"
                    st.session_state["mail_msg"]    = mail_err or "Unknown error"
            st.rerun()

        with st.expander("📋 AI Content"):
            st.json(_ai)
        with st.expander("📊 Career Data"):
            st.write(f"**Roles generated:** {len(_rows)}")
            st.write(f"**Domains:** {_dmap}")

# ════════════════════════════════════════════════════════
# TAB 3 — SPECIAL CONSULTATION
# ════════════════════════════════════════════════════════
with tab3:
    st.subheader("Special Consultation Note")
    st.markdown(
        '<p style="color:#555;font-size:15px;margin-bottom:24px;">'
        "Type your custom consultation note below. This will appear directly under the "
        "<strong>Prescription</strong> heading in the generated PDF and Word document, "
        "followed by the AI-generated domain-specific content as usual."
        "</p>",
        unsafe_allow_html=True
    )

    st.markdown(
        '<div class="card" style="margin-bottom:12px;">'
        '<div class="card-label">Example</div>'
        '<div class="card-text" style="color:#444;font-style:italic;">'
        '"As per our discussion, you completed B.Com and have two years of experience in accounts. '
        'You cannot directly become a Data Analyst at this stage, but this programme will bridge that gap '
        'and help you transition into a data-driven role with the right skills and projects."'
        "</div></div>",
        unsafe_allow_html=True
    )

    consultation_input = st.text_area(
        "Consultation Note *",
        value=st.session_state.get("consultation_note", ""),
        height=160,
        placeholder=(
            "e.g. As per our discussion, you completed B.Com and have two years of experience in accounts. "
            "You cannot directly become a Data Analyst at this stage, but this programme will bridge that gap "
            "and help you transition into a data-driven role with the right skills and projects."
        ),
        key="consultation_note_input"
    )

    col_save, col_clear, _ = st.columns([2, 2, 5])
    with col_save:
        if st.button("💾 Save Note", type="primary", key="save_consult_note"):
            note_val = consultation_input.strip()
            if note_val:
                # Capitalise first letter, ensure ends with period
                note_val = note_val[0].upper() + note_val[1:]
                if note_val[-1] not in ".!?":
                    note_val += "."
            st.session_state["consultation_note"] = note_val
            if note_val:
                st.success("✅ Consultation note saved. Go to the Application tab and generate the prescription — your note will appear under the Prescription heading.")
            else:
                st.info("ℹ️ Note cleared. The Prescription section will use the AI-generated intro line only.")

    with col_clear:
        if st.button("🗑️ Clear Note", key="clear_consult_note"):
            st.session_state["consultation_note"] = ""
            st.info("ℹ️ Consultation note cleared.")

    # Show current saved note
    saved_note = st.session_state.get("consultation_note", "")
    if saved_note:
        st.markdown("---")
        st.markdown("**Currently saved note (will appear in next generated PDF):**")
        st.markdown(
            f'<div class="card"><div class="card-text">{saved_note}</div></div>',
            unsafe_allow_html=True
        )
